# -*- coding: utf-8 -*-
"""Genera el documento Word con el esquema y flujo de trabajo de marketing-map."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

INDIGO = RGBColor(0x4F, 0x46, 0xE5)
DARK = RGBColor(0x0F, 0x17, 0x2A)
GRAY = RGBColor(0x47, 0x55, 0x69)

doc = Document()

# Estilo base
normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(11)
normal.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

def h1(text):
    p = doc.add_heading(level=1)
    r = p.add_run(text); r.font.color.rgb = INDIGO; r.font.size = Pt(16); r.bold = True
    return p

def h2(text):
    p = doc.add_heading(level=2)
    r = p.add_run(text); r.font.color.rgb = DARK; r.font.size = Pt(13); r.bold = True
    return p

def para(text, italic=False, color=None, size=11):
    p = doc.add_paragraph()
    r = p.add_run(text); r.italic = italic; r.font.size = Pt(size)
    if color: r.font.color.rgb = color
    return p

def bullet(text):
    doc.add_paragraph(text, style='List Bullet')

def numbered(text):
    doc.add_paragraph(text, style='List Number')

def table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].text = ''
        run = hdr[i].paragraphs[0].add_run(htext); run.bold = True; run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ''
            run = cells[i].paragraphs[0].add_run(str(val)); run.font.size = Pt(10)
    return t

# ----------------- Portada -----------------
title = doc.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('marketing-map'); r.bold = True; r.font.size = Pt(34); r.font.color.rgb = INDIGO
sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('Esquema y Flujo de Trabajo'); r.font.size = Pt(18); r.font.color.rgb = DARK
sub2 = doc.add_paragraph(); sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub2.add_run('Sistema de captación, clasificación y seguimiento de clientes con Inteligencia Artificial')
r.italic = True; r.font.size = Pt(11); r.font.color.rgb = GRAY
meta = doc.add_paragraph(); meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run('Documento de trabajo · Junio 2026'); r.font.size = Pt(10); r.font.color.rgb = GRAY
doc.add_paragraph()

# ----------------- 1. Resumen ejecutivo -----------------
h1('1. Resumen ejecutivo')
para('marketing-map es un sistema omnicanal que capta clientes potenciales (leads), los clasifica '
     'automáticamente con Inteligencia Artificial, los guarda en una base de datos central (CRM) y les '
     'da seguimiento automático hasta convertirlos en clientes. Todo el tratamiento de datos se realiza '
     'con consentimiento y conforme a la Ley Orgánica de Protección de Datos Personales (LOPDP) de Ecuador.')
para('Principio rector: atraer y convertir con herramientas oficiales y datos consentidos. No se realiza '
     'vigilancia ni perfilado de personas en redes sociales (práctica ilegal y sancionable).',
     italic=True, color=GRAY)

# ----------------- 2. Esquema de trabajo (componentes) -----------------
h1('2. Esquema de trabajo (arquitectura)')
para('El sistema se compone de cinco piezas que trabajan de forma integrada:')
table(['Componente', 'Función'],
      [['Landing page', 'Página pública donde el cliente deja sus datos. Alimenta el CRM.'],
       ['Backend + IA (cerebro)', 'Recibe los datos, clasifica con IA (Claude), responde y dispara el seguimiento.'],
       ['Base de datos (CRM)', 'Guarda contactos, conversaciones, productos, secuencias y consentimientos.'],
       ['Panel de control', 'Donde el equipo ve clientes, conversaciones, métricas y gestiona todo.'],
       ['Canales', 'WhatsApp, Instagram/Facebook, Email y Web (se conectan según se requiera).']])
doc.add_paragraph()

h2('Stack tecnológico')
table(['Pieza', 'Tecnología'],
      [['Base de datos y autenticación', 'Supabase (PostgreSQL)'],
       ['Backend / motor', 'Node.js + TypeScript'],
       ['Inteligencia Artificial', 'Claude (Anthropic)'],
       ['Panel de control', 'Next.js + React'],
       ['Infraestructura (hosting)', 'Coolify'],
       ['Canales', 'WhatsApp Cloud API, Meta Graph, Email']])

# ----------------- 3. Flujo de trabajo del CLIENTE -----------------
doc.add_page_break()
h1('3. Flujo de trabajo del cliente (ciclo de vida del lead)')
para('Así avanza cada persona dentro del sistema, de forma automática:')
numbered('CAPTACIÓN — El cliente llega por un anuncio, contenido o referido y deja sus datos en la '
         'landing (o entra por WhatsApp, o se importa desde una base existente).')
numbered('REGISTRO — Se crea/actualiza el contacto en el CRM, sin duplicados, con su canal de origen.')
numbered('CLASIFICACIÓN CON IA — La IA analiza el mensaje y asigna: etapa del embudo, nivel de interés, '
         'score de compra (0–100), intención y un resumen del cliente.')
numbered('RESPUESTA / ASIGNACIÓN — La IA responde como asesor (en WhatsApp) o el lead queda listo para '
         'que el equipo lo atienda, según el canal.')
numbered('SEGUIMIENTO AUTOMÁTICO — Si el cliente no responde, una secuencia le envía recordatorios y '
         'ofertas en los tiempos definidos, respetando el horario y la baja.')
numbered('CONVERSIÓN — Cuando el cliente compra, se marca como “Cliente”. El equipo cierra la venta.')
numbered('BAJA (opt-out) — Si el cliente pide no recibir más mensajes (“STOP/BAJA”), se da de baja '
         'automáticamente y deja de recibir comunicaciones. Queda registrado (auditable).')
doc.add_paragraph()
para('Embudo de etapas:', color=DARK)
para('Nuevo  →  Interesado  →  Calificado  →  Negociando  →  Cliente   (o  →  Perdido)', color=INDIGO)

# ----------------- 4. Flujo de trabajo del EQUIPO -----------------
h1('4. Flujo de trabajo del equipo (operación diaria)')
numbered('Iniciar sesión en el panel (acceso restringido por usuario y contraseña).')
numbered('Revisar el TABLERO: nuevos leads del día, embudo por etapa y métricas.')
numbered('Atender CLIENTES priorizando por score (mayor probabilidad de compra primero).')
numbered('Responder/continuar conversaciones; la IA ya dejó respuestas y clasificación.')
numbered('Avanzar la etapa de cada cliente o eliminar los registros incompletos.')
numbered('Cargar/editar PRODUCTOS y precios para que la IA venda con información correcta.')
numbered('Importar nuevas bases de contactos (con consentimiento) cuando corresponda.')
numbered('Exportar AUDIENCIAS para campañas de anuncios (Meta/Google) y atraer más leads.')

# ----------------- 5. Roles -----------------
h1('5. Roles y responsabilidades')
table(['Rol', 'Responsabilidades'],
      [['Administrador', 'Configura el sistema, productos, secuencias y usuarios. Acceso total.'],
       ['Socia / Asesor', 'Atiende clientes, da seguimiento, cierra ventas y actualiza etapas.'],
       ['Sistema (IA)', 'Clasifica, responde, da seguimiento y aplica las bajas automáticamente.']])

# ----------------- 6. Cumplimiento legal -----------------
h1('6. Cumplimiento legal (LOPDP)')
bullet('Consentimiento: se solicita autorización explícita antes de tratar los datos (landing e importación).')
bullet('Baja garantizada: todo cliente puede darse de baja en cualquier momento y se respeta al instante.')
bullet('Datos mínimos: se guarda solo lo necesario para contactar y vender.')
bullet('Sin vigilancia: no se rastrea ni perfila a personas en redes sociales.')
bullet('Política de privacidad publicada en la landing.')

# ----------------- 7. Fases -----------------
doc.add_page_break()
h1('7. Fases del proyecto')
table(['Fase', 'Alcance', 'Estado'],
      [['1 · Base', 'CRM, IA, landing, seguimiento, panel, login, importador, audiencias', 'Implementado'],
       ['2 · WhatsApp', 'Conexión de WhatsApp Cloud API (respuestas automáticas en vivo)', 'En proceso'],
       ['3 · Omnicanal', 'Instagram, Facebook y Email integrados', 'Pendiente'],
       ['4 · Inteligencia', 'Panel de tendencias agregadas y A/B testing de mensajes', 'Pendiente'],
       ['5 · Escala', 'Roles avanzados, reportes y automatización de anuncios', 'Pendiente']])

# ----------------- 8. KPIs -----------------
h1('8. Indicadores clave (KPIs) a monitorear')
bullet('Leads nuevos por día / semana.')
bullet('Tasa de conversión (leads que se vuelven clientes).')
bullet('Tiempo de respuesta al cliente.')
bullet('Distribución del embudo por etapa.')
bullet('Bajas (opt-out) y su motivo.')
bullet('Costo por lead y retorno de las campañas de anuncios.')

# ----------------- 9. Accesos -----------------
h1('9. Accesos del sistema')
table(['Recurso', 'Enlace'],
      [['Panel de control', 'https://<tu-dominio-panel>'],
       ['Landing (clientes)', 'https://<tu-dominio-landing>'],
       ['Backend (API)', 'https://<tu-dominio-backend>']])
para('Nota de seguridad: las credenciales compartidas durante la configuración deben rotarse, y la '
     'contraseña del panel debe cambiarse periódicamente.', italic=True, color=GRAY)

out = r'C:\Users\mapos\Dropbox\Programas\Marketing_MAP\Esquema_de_Trabajo_marketing-map.docx'
doc.save(out)
print('Documento generado:', out)
